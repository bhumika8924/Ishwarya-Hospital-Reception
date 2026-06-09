from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
SOURCE_MD = PROJECT_ROOT / "docs" / "project_documentation_for_beginners.md"
OUTPUT_DOCX = PROJECT_ROOT / "docs" / "Hospital_Reception_Monitor_Beginner_Documentation.docx"


HEADING_BLUE = RGBColor(46, 116, 181)
HEADING_DARK_BLUE = RGBColor(31, 77, 120)
MUTED_GRAY = RGBColor(89, 89, 89)
CODE_FILL = "F4F6F8"
CODE_BORDER = "D9E2EC"
TABLE_HEADER_FILL = "E8EEF5"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.bold = bold
    paragraph.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_paragraph_border(paragraph, color=CODE_BORDER, fill=CODE_FILL):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)

    for side in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "4")
        element.set(qn("w:color"), color)

    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_code_block(document, lines):
    text = "\n".join(lines).rstrip()
    paragraph = document.add_paragraph()
    paragraph.style = "Code Block"
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.12)
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(8)
    set_paragraph_border(paragraph)
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(32, 32, 32)


def add_inline_markdown(paragraph, text, bold=False):
    pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")
    position = 0
    for match in pattern.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position:match.start()])
            run.bold = bold
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(32, 32, 32)
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        run.bold = bold


def configure_document(document):
    section = document.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color, before, after in [
        ("Heading 1", 16, HEADING_BLUE, 14, 8),
        ("Heading 2", 13, HEADING_BLUE, 10, 5),
        ("Heading 3", 12, HEADING_DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    code_style = styles.add_style("Code Block", 1)
    code_style.font.name = "Consolas"
    code_style.font.size = Pt(9)
    code_style.paragraph_format.line_spacing = 1.05


def add_cover(document):
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(130)
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("Hospital Reception Monitor")
    run.font.name = "Calibri"
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = HEADING_DARK_BLUE

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(26)
    run = subtitle.add_run("Beginner-Friendly Project Documentation")
    run.font.name = "Calibri"
    run.font.size = Pt(16)
    run.font.color.rgb = MUTED_GRAY

    summary = document.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary.paragraph_format.left_indent = Inches(0.5)
    summary.paragraph_format.right_indent = Inches(0.5)
    summary.paragraph_format.space_after = Pt(18)
    add_inline_markdown(
        summary,
        "A simple explanation of the project problem, solution approach, implementation flow, setup steps, and future improvements.",
    )

    table = document.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        row.cells[0].width = Inches(1.8)
        row.cells[1].width = Inches(4.3)
    data = [
        ("Project", "Hospital Reception Monitor"),
        ("Purpose", "Receptionist count and visitor entry/exit count from video"),
        ("Main Tools", "Streamlit, OpenCV, YOLOv8, Python"),
        ("Audience", "Managers, reviewers, and beginners learning the project"),
    ]
    for row, (label, value) in zip(table.rows, data):
        set_cell_shading(row.cells[0], TABLE_HEADER_FILL)
        set_cell_text(row.cells[0], label, bold=True)
        set_cell_text(row.cells[1], value)

    document.add_section(WD_SECTION.NEW_PAGE)


def add_footer(document):
    for section in document.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("Hospital Reception Monitor Documentation")
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED_GRAY


def add_overview_table(document):
    document.add_heading("Quick Project Snapshot", level=1)
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(1.9)
    table.columns[1].width = Inches(4.4)
    set_cell_shading(table.rows[0].cells[0], TABLE_HEADER_FILL)
    set_cell_shading(table.rows[0].cells[1], TABLE_HEADER_FILL)
    set_cell_text(table.rows[0].cells[0], "Item", bold=True)
    set_cell_text(table.rows[0].cells[1], "Explanation", bold=True)
    rows = [
        ("Problem", "Manual CCTV/video counting is slow and can become inaccurate."),
        ("Solution", "Upload a video and automatically count receptionists, visitor entries, and visitor exits."),
        ("Detection", "YOLOv8 detects people. Tracking IDs follow the same person across frames."),
        ("Receptionist Logic", "Reception zone plus uniform/color checks confirm staff members."),
        ("Visitor Logic", "Two saved lines identify entry and exit direction."),
    ]
    for label, value in rows:
        row = table.add_row()
        set_cell_text(row.cells[0], label, bold=True)
        set_cell_text(row.cells[1], value)
    document.add_paragraph()


def markdown_to_docx(document, markdown):
    in_code = False
    code_lines = []
    pending_list = []

    def flush_list():
        nonlocal pending_list
        pending_list = []

    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()
        if line.startswith("```"):
            if in_code:
                add_code_block(document, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if not line.strip():
            flush_list()
            continue

        if line.startswith("# "):
            continue
        if line.startswith("## "):
            document.add_heading(line[3:].strip(), level=1)
            continue
        if line.startswith("### "):
            document.add_heading(line[4:].strip(), level=2)
            continue

        bullet_match = re.match(r"^- (.+)", line)
        number_match = re.match(r"^\d+\. (.+)", line)
        if bullet_match:
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline_markdown(paragraph, bullet_match.group(1))
            continue
        if number_match:
            paragraph = document.add_paragraph(style="List Number")
            add_inline_markdown(paragraph, number_match.group(1))
            continue

        paragraph = document.add_paragraph()
        add_inline_markdown(paragraph, line)

    if in_code and code_lines:
        add_code_block(document, code_lines)


def main():
    markdown = SOURCE_MD.read_text(encoding="utf-8")
    document = Document()
    configure_document(document)
    add_cover(document)
    add_footer(document)
    add_overview_table(document)
    markdown_to_docx(document, markdown)
    document.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
