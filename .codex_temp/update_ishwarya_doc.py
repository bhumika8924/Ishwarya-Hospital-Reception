from docx import Document


DOC_PATH = "D:/Bhumi/Team Computers/Ishwaraya_Hospital_Reception-main/.codex_temp/Ishwarya Reception.docx"


def set_text(paragraph, text):
    paragraph.clear()
    paragraph.add_run(text)


def replace_exact(paragraphs, old, new):
    for paragraph in paragraphs:
        if paragraph.text == old:
            set_text(paragraph, new)
            return
    raise ValueError(f"Could not find paragraph: {old!r}")


document = Document(DOC_PATH)
paragraphs = document.paragraphs

replacements = {
    "There are two receptionist detection modes:": (
        "There is one combined receptionist dress-code detection flow:"
    ),
    "Pink/reference-uniform mode.": (
        "Pink/reference-uniform checking using the reference image and lanyard/id-card color."
    ),
    "Blue-saree mode.": (
        "Blue-saree checking using direct blue-color detection in the body area."
    ),
    "4. Why There Are Two Apps": "4. Why The App Is Combined",
    "The project has two separate Streamlit apps because receptionist uniform detection is different in two cases.": (
        "The project now uses one Streamlit app because pink/reference uniform and blue saree are both valid receptionist dress-code rules."
    ),
    "Pink/reference-uniform app": "Combined dress-code app",
    "This mode uses a reference image:": (
        "This app uses a reference image for pink/reference-uniform checking:"
    ),
    "The app compares the detected person's clothing color with the reference image. It also checks for red or pink lanyard/id-card color near the neck area.": (
        "The app compares the detected person's clothing color with the reference image, checks for red or pink lanyard/id-card color near the neck area, and also checks blue saree color directly from the detected person's body area."
    ),
    "Blue-saree app": "Blue-saree compatibility launcher",
    "This mode does not use a reference image. Instead, it directly checks whether enough blue color is present in the detected person's body area.": (
        "The old blue-saree launcher still works, but it opens the same combined app."
    ),
    "This was added because a blue saree uniform needs a different detection rule from the pink/reference uniform.": (
        "The blue saree rule is now inside the same core analysis file, so both dress-code checks are maintained together."
    ),
    "`assets/receptionist_uniform_ref.png`": "`assets/receptionist_uniform_ref.png`",
    "This is the reference uniform image for the pink/reference-uniform app.": (
        "This is the reference uniform image for the pink/reference-uniform check."
    ),
    "Step 3: Run pink/reference-uniform app": "Step 3: Run the combined dress-code app",
    "Step 4: Run blue-saree app": "Step 4: Optional old blue-saree shortcut",
    "The main analysis call for the pink/reference app is:": (
        "The main analysis call for the combined dress-code app is:"
    ),
    "The blue-saree app calls a similar function but uses min_blue_pixels instead of min_red.": (
        "The same analysis call now accepts both min_red and min_blue_pixels, so pink/reference and blue-saree dress-code checks run together."
    ),
    "Pink/reference-uniform check": "Pink/reference-uniform check",
    "For the pink/reference app:": "For the pink/reference check:",
    "So in the pink/reference app, a person can be confirmed as receptionist if:": (
        "For the combined app, a person can be confirmed as receptionist if:"
    ),
    "The uniform color matches or lanyard/id-card color is detected.": (
        "The pink/reference uniform color matches, lanyard/id-card color is detected, or blue-saree color is detected."
    ),
    "For the blue-saree app:": "For the blue-saree check:",
    "For the pink/reference app, replace this image if the uniform changes:": (
        "For the pink/reference check, replace this image if the uniform changes:"
    ),
    "For the blue-saree app, the blue HSV values are hardcoded in:": (
        "For the blue-saree check, the blue HSV values are now hardcoded in:"
    ),
    "core/counter_core_blue.py": "core/counter_core.py",
    "Lanyard/id-card color in pink/reference mode.": (
        "Lanyard/id-card color for the pink/reference rule."
    ),
    "Blue-saree color in blue mode.": "Blue-saree color for the blue dress-code rule.",
    "Run pink/reference-uniform app:": "Run combined dress-code app:",
    "Run blue-saree app:": "Optional old blue-saree shortcut:",
}

for old, new in replacements.items():
    replace_exact(paragraphs, old, new)

folder_structure = """Ishwaraya_Hospital_Reception-main/
  apps/
    streamlit_app.py
    streamlit_app_blue.py

  core/
    counter_core.py
    counter_core_blue.py

  tools/
    two_line_visitor_counter.py

  config/
    reception_zone.json
    two_line_counter_lines.json

  assets/
    yolov8n.pt
    receptionist_uniform_ref.png

  data/
    videos/

  docs/
    use_case.md
    project_documentation_for_beginners.md

  .streamlit/
    config.toml

  requirements.txt
  README.md
  streamlit_app.py
  streamlit_app_blue.py"""

folder_structure_new = """Ishwaraya_Hospital_Reception-main/
  apps/
    streamlit_app.py              <- Combined dress-code receptionist counter
    streamlit_app_blue.py         <- Old shortcut that opens the combined app

  core/
    counter_core.py               <- Combined pink/reference and blue-saree analysis logic
    counter_core_blue.py          <- Old import shortcut for the combined core

  tools/
    two_line_visitor_counter.py

  config/
    reception_zone.json
    two_line_counter_lines.json

  assets/
    yolov8n.pt
    receptionist_uniform_ref.png

  data/
    videos/

  docs/
    use_case.md
    project_documentation_for_beginners.md

  .streamlit/
    config.toml

  requirements.txt
  README.md
  streamlit_app.py
  streamlit_app_blue.py"""

replace_exact(paragraphs, folder_structure, folder_structure_new)

analysis_call_old = """result = analyze_video(
    video_path=video_path,
    model_path=MODEL_PATH,
    reference_path=REFERENCE_PATH,
    zone_path=ZONE_PATH,
    two_line_path=two_line_path,
    ref_w=ref_w,
    ref_h=ref_h,
    entry_order_option=entry_order_option,
    yolo_conf=yolo_conf,
    yolo_iou=yolo_iou,
    match_threshold=match_threshold,
    min_red=min_red,
    min_overlap=min_overlap,
    analyze_every=analyze_every,
    on_progress=update_progress,
    on_frame=update_frame,
    preview_every=5,
)"""

analysis_call_new = """result = analyze_video(
    video_path=video_path,
    model_path=MODEL_PATH,
    reference_path=REFERENCE_PATH,
    zone_path=ZONE_PATH,
    two_line_path=two_line_path,
    ref_w=ref_w,
    ref_h=ref_h,
    entry_order_option=entry_order_option,
    yolo_conf=yolo_conf,
    yolo_iou=yolo_iou,
    match_threshold=match_threshold,
    min_red=min_red,
    blue_match_threshold=blue_match_threshold,
    min_blue_pixels=min_blue_pixels,
    min_overlap=min_overlap,
    analyze_every=analyze_every,
    on_progress=update_progress,
    on_frame=update_frame,
    preview_every=5,
)"""

replace_exact(paragraphs, analysis_call_old, analysis_call_new)

document.save(DOC_PATH)
